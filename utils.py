"""
TalentSync AI — Utility Functions v3.0
File text extraction (pdfplumber → pypdf → fallback), response builders, helpers.
"""

import io
import re
from typing import Dict, Any, Optional


def extract_text_from_file(content: bytes, filename: str) -> str:
    """
    Extract plain text from uploaded file bytes.
    Supports: .txt, .pdf (multi-column aware), .docx
    """
    name = (filename or "").lower()

    if name.endswith(".txt") or name.endswith(".md"):
        return content.decode("utf-8", errors="replace").strip()

    if name.endswith(".pdf"):
        return _extract_pdf(content)

    if name.endswith(".docx"):
        return _extract_docx(content)

    # Fallback — try plain text decode
    try:
        decoded = content.decode("utf-8", errors="replace").strip()
        if len(decoded) > 50:
            return decoded
    except Exception:
        pass

    raise ValueError(f"Unsupported file type: {filename}. Please upload PDF, TXT, or DOCX.")


def _extract_pdf(content: bytes) -> str:
    """
    Extract text from PDF.
    Strategy: pdfplumber (multi-column, tables) → pypdf → raw fallback.
    """
    # Strategy 1: pdfplumber — handles multi-column and tabular layouts well
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            pages = []
            for page in pdf.pages:
                # Extract regular text
                text = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
                # Also extract tables and append as plain text
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        row_text = " | ".join(str(cell or "").strip() for cell in row if cell)
                        if row_text.strip():
                            text += "\n" + row_text
                pages.append(text)
            full_text = "\n".join(pages).strip()
            if len(full_text) > 50:
                return full_text
    except ImportError:
        pass
    except Exception:
        pass

    # Strategy 2: pypdf — simpler but works for most standard PDFs
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(content))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        if len(text.strip()) > 50:
            return text.strip()
    except ImportError:
        pass
    except Exception:
        pass

    # Strategy 3: pdfminer (if installed)
    try:
        from pdfminer.high_level import extract_text_to_fp
        from pdfminer.layout import LAParams
        output = io.StringIO()
        extract_text_to_fp(io.BytesIO(content), output, laparams=LAParams())
        text = output.getvalue().strip()
        if len(text) > 50:
            return text
    except ImportError:
        pass
    except Exception:
        pass

    # Strategy 4: raw stream extraction (last resort)
    raw = content.decode("latin-1", errors="replace")
    streams = re.findall(r"stream(.*?)endstream", raw, re.DOTALL)
    extracted = " ".join(streams)
    clean = re.sub(r"[^\x20-\x7E\n]", " ", extracted)
    clean = re.sub(r"\s+", " ", clean).strip()

    if len(clean) < 50:
        raise ValueError(
            "Could not extract text from this PDF. "
            "It may be image-only (scanned). Please paste the resume text instead."
        )
    return clean


def _extract_docx(content: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs).strip()
    except ImportError:
        pass

    # Fallback: unzip and parse XML
    import zipfile
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as z:
            with z.open("word/document.xml") as f:
                xml = f.read().decode("utf-8", errors="replace")
        text = re.sub(r"<[^>]+>", " ", xml)
        text = re.sub(r"\s+", " ", text).strip()
        return text
    except Exception:
        raise ValueError("Could not extract text from DOCX. Please paste the text manually.")


def build_pipeline_response(
    candidate_id: str,
    parsed: Dict[str, Any],
    normalized_skills: list,
    match_data: Dict[str, Any],
    logs: list,
    pipeline_time_ms: int,
    emerging_skills: list = None,
) -> Dict[str, Any]:
    """Build the standard pipeline response envelope."""
    return {
        "candidate_id": candidate_id,
        "parsed": parsed,
        "normalized_skills": normalized_skills,
        "match": match_data,
        "pipeline_logs": logs,
        "pipeline_time_ms": pipeline_time_ms,
        "emerging_skills_detected": emerging_skills or [],
        "summary": {
            "name": parsed.get("name", "Unknown"),
            "title": parsed.get("title", ""),
            "skills_count": len(normalized_skills),
            "match_score": match_data.get("match_score", 0),
            "verdict": match_data.get("verdict", ""),
            "hiring_signal": match_data.get("hiring_signal", "consider"),
            "vector_similarity_score": match_data.get("vector_similarity_score"),
        }
    }


def sanitize_text(text: str, max_chars: int = 8000) -> str:
    """Sanitize and truncate resume text before sending to AI."""
    clean = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", text)
    clean = re.sub(r"\s+", " ", clean).strip()
    if len(clean) > max_chars:
        clean = clean[:max_chars] + "\n[...truncated for processing...]"
    return clean


def compute_quality_score(parsed: Dict[str, Any]) -> int:
    """Compute a 0-100 quality score for a parsed resume."""
    score = 0
    checks = [
        (bool(parsed.get("name")), 15),
        (bool(parsed.get("email")), 10),
        (bool(parsed.get("title")), 10),
        (len(parsed.get("skills", [])) >= 5, 20),
        (len(parsed.get("experience", [])) >= 1, 20),
        (len(parsed.get("education", [])) >= 1, 15),
        (bool(parsed.get("summary")), 10),
    ]
    for condition, points in checks:
        if condition:
            score += points
    return score
